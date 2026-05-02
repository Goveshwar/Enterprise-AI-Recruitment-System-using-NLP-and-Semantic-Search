# import os
# import re
# import json
# import logging
# import magic
# import fitz
# import pytesseract
# import spacy
# from datetime import datetime
# from docx import Document
# from PIL import Image, ImageOps
# from sentence_transformers import SentenceTransformer, util
# from company_extractor import CompanyExtractor
# from skill_matcher import SkillMatcher

# # --- SYSTEM CONFIG ---
# logging.basicConfig(level=logging.INFO, format='%(message)s')
# logger = logging.getLogger(__name__)

# class UniversalDynamicParser:
#     def __init__(self, jd_path=None):
#         logger.info("Initializing Universal AI Engine...")
#         pytesseract.pytesseract.tesseract_cmd = r'/usr/bin/tesseract'
#         self.embedder = SentenceTransformer('all-MiniLM-L6-v2')
#         self.nlp = spacy.load("en_core_web_md")
#         self.company_extractor = CompanyExtractor()
#         self.skill_matcher = SkillMatcher()
        
#         # Load JD for Dynamic Semantic Matching
#         self.jd_text = self.extract_raw_text(jd_path) if jd_path else ""
#         self.jd_vector = self.embedder.encode(self.jd_text) if self.jd_text else None

#     def extract_raw_text(self, path):
#         """Universal Ingestion: Detects file type via Magic Bytes, not extensions."""
#         if not path or not os.path.exists(path): return ""
#         mime = magic.Magic(mime=True).from_file(path)
#         try:
#             if "image" in mime:
#                 with Image.open(path) as img:
#                     img = ImageOps.autocontrast(img.convert('L'))
#                     img = img.resize((img.width * 2, img.height * 2), Image.LANCZOS)
#                     # PSM 1: Auto layout analysis for multi-column resumes
#                     return pytesseract.image_to_string(img, config='--oem 3 --psm 1')
#             if "pdf" in mime:
#                 text = ""
#                 with fitz.open(path) as doc:
#                     for page in doc: text += page.get_text()
#                     if len(text.strip()) < 100: # Scanned PDF Fallback
#                         for page in doc:
#                             pix = page.get_pixmap(dpi=300)
#                             img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
#                             text += pytesseract.image_to_string(img, config='--psm 1')
#                 return text
#             if "officedocument" in mime:
#                 doc = Document(path)
#                 return "\n".join([p.text for p in doc.paragraphs])
#             return open(path, 'r', errors='ignore').read()
#         except Exception as e:
#             logger.error(f"Ingestion Fail: {e}")
#             return ""

#     def _scrub(self, text):
#         """Cleans artifacts, city/state locations, and sidebar bleed."""
#         if not text: return ""
#         # Remove locations (City, ST or Country)
#         loc_patterns = r'(\b[A-Z][a-z]+, [A-Z]{2}\b|Mountain View|Redmond|Palo Alto|San Francisco|Remote|India|USA)'
#         text = re.sub(loc_patterns, '', text)
#         # Remove common artifacts
#         noise = [r"\|", r"•", r"Experience", r"Skills", r"9 ", r"Page \d"]
#         for p in noise:
#             text = re.sub(p, '', text, flags=re.IGNORECASE)
#         return " ".join(text.split()).strip("@, -")

#     def parse(self, resume_path):
#         raw_text = self.extract_raw_text(resume_path)
#         if not raw_text.strip(): 
#             return {"error": "Empty File"}

#         bio = self._extract_verified_bio(raw_text)
#         experience = self._extract_experience_universal(raw_text)
#         education = self._extract_education_universal(raw_text)

#         # ✅ NEW: Skill Matching
#         skill_analysis = self.skill_matcher.match_skills(
#             self.jd_text,
#             raw_text
#         )

#         return {
#             "biographical": bio,
#             "experience": experience,
#             "education": education,
#             "skills": self._extract_skills(raw_text),
#             "skill_analysis": skill_analysis,   # 👈 NEW
#             "dynamic_match": self._calculate_match(raw_text),
#             "metadata": {
#                 "source_file": os.path.basename(resume_path),
#                 "processed_at": datetime.now().isoformat()
#             }
#         }

#     def _extract_verified_bio(self, text):
#         email = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
#         phone = re.search(r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text)
        
#         # Blacklist common words mislabeled as NAMES in tech if resumes
#         blacklist = ["Java", "Python", "SQL", "Azure", "AWS", "Engineer", "Developer", "Resume"]
        
#         doc = self.nlp(text[:500])
#         potential_names = [ent.text for ent in doc.ents if ent.label_ == "PERSON"]
#         valid_name = "Unknown"
#         for n in potential_names:
#             if not any(b in n for b in blacklist) and len(n.split()) > 1:
#                 valid_name = n
#                 break
        
#         if valid_name == "Unknown": # Fallback to first line
#             valid_name = text.split('\n')[0]
            
#         return {"name": self._scrub(valid_name), "email": email.group(0) if email else None, "phone": phone.group(0) if phone else None}

#     def _extract_experience_universal(self, text):
#         """Finds dates and looks for nearest Company/Role entities."""
#         jobs = []
#         date_pattern = r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|20\d{2}|\d{4}).*?(?:Present|Current|20\d{2}|\d{4}))'
        
#         # We split the text into chunks based on where dates are located
#         chunks = re.split(date_pattern, text)
#         if len(chunks) < 2: return []

#         for i in range(1, len(chunks)-1, 2):
#             tenure = chunks[i].strip()
#             body = chunks[i+1].strip()
            
#             # Extract Company using NLP ORG label within this specific block
#             doc = self.nlp(body[:200])
#             org = next((ent.text for ent in doc.ents if ent.label_ == "ORG"), "Company")
            
#             # Extract Role: It's usually the first line that isn't the ORG
#             lines = [l.strip() for l in body.split('\n') if len(l.strip()) > 3]
#             role = lines[0] if lines else "Professional"
#             if org in role and len(lines) > 1: role = lines[1]

#             jobs.append({
#                 "employer": self._scrub(org),
#                 "role": self._scrub(role),
#                 "tenure": tenure
#             })
#         return jobs[:5]

#     def _extract_education_universal(self, text):
#         # Broad patterns for degrees
#         degrees = re.findall(r'(Bachelor|Master|B\.S|M\.S|B\.Tech|MBA|PhD|B\.E|Graduate)', text, re.I)
#         # Search for universities in the text
#         orgs = [ent.text for ent in self.nlp(text).ents if ent.label_ == "ORG" and "University" in ent.text]
#         return [{"school": orgs[0] if orgs else "University", "degree": degrees[0] if degrees else "Degree"}]

#     def _extract_skills(self, text):
#         # Comprehensive Tech dictionary
#         bank = ["Python", "SQL", "Java", "Azure", "AWS", "Docker", "ETL", "Spark", "Hadoop", "Kubernetes", "Airflow"]
#         return [s for s in bank if s.lower() in text.lower()]

#     def _calculate_match(self, text):
#         if self.jd_vector is None: return "0%"
#         score = util.cos_sim(self.jd_vector, self.embedder.encode(text)).item()
#         return f"{round(score * 100, 2)}%"

# # --- EXECUTION ---
# if __name__ == "__main__":
#     # The system will now correctly process ANY JD and ANY Resume layout
#     JD_PATH = "/home/credentek/Downloads/software-engineer-job-description-D13502.png"
#     RES_PATH = "resume2.docx"

#     parser = UniversalDynamicParser(jd_path=JD_PATH)
#     print(json.dumps(parser.parse(RES_PATH), indent=4))




#########################



# import os
# import re
# import logging
# import magic
# import fitz
# import pytesseract
# import spacy
# from datetime import datetime
# from docx import Document
# from PIL import Image, ImageOps

# from transformers import pipeline
# from skill_matcher import SkillMatcher

# logging.basicConfig(level=logging.INFO, format="%(message)s")
# logger = logging.getLogger(__name__)


# class UniversalATSParser:

#     def __init__(self, jd_path=None):

#         logger.info("Initializing ATS Engine...")

#         pytesseract.pytesseract.tesseract_cmd = "/usr/bin/tesseract"

#         self.nlp = spacy.load("en_core_web_sm")

#         self.ner = pipeline(
#             "token-classification",
#             model="dslim/bert-base-NER",
#             aggregation_strategy="simple"
#         )

#         self.skill_matcher = SkillMatcher()

#         self.jd_path = jd_path
#         self.jd_text = self._load_jd_text(jd_path)

#         logger.info("ATS Engine Ready ✅")

#     # -------------------------
#     # LOAD JD
#     # -------------------------
#     def _load_jd_text(self, jd_path):
#         if not jd_path or not os.path.exists(jd_path):
#             return ""
#         return self.extract_text(jd_path)

#     # -------------------------
#     # UNIVERSAL TEXT EXTRACTION
#     # -------------------------
#     def extract_text(self, path):

#         if not path or not os.path.exists(path):
#             return ""

#         try:
#             mime = magic.Magic(mime=True).from_file(path)

#             # IMAGE
#             if "image" in mime:
#                 img = Image.open(path)
#                 img = ImageOps.autocontrast(img.convert("L"))
#                 return pytesseract.image_to_string(img, config="--psm 6")

#             # PDF
#             if "pdf" in mime:
#                 text = ""
#                 doc = fitz.open(path)

#                 for page in doc:
#                     text += page.get_text()

#                 if len(text.strip()) < 100:
#                     for page in doc:
#                         pix = page.get_pixmap(dpi=300)
#                         img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
#                         text += pytesseract.image_to_string(img, config="--psm 6")

#                 doc.close()
#                 return text

#             # DOCX
#             if "officedocument" in mime:
#                 doc = Document(path)
#                 return "\n".join(p.text for p in doc.paragraphs)

#             # TXT
#             with open(path, "r", errors="ignore") as f:
#                 return f.read()

#         except Exception as e:
#             logger.error(f"File read error: {e}")
#             return ""

#     # -------------------------
#     # BIO
#     # -------------------------
#     def extract_bio(self, text):

#         email = re.search(r"[\w\.-]+@[\w\.-]+\.\w+", text)
#         phone = re.search(r"(\+?\d{1,3}[-.\s]?)?\d{10}", text)

#         doc = self.nlp(text[:1200])

#         name = "Unknown"

#         for ent in doc.ents:
#             if ent.label_ == "PERSON":
#                 if 1 <= len(ent.text.split()) <= 3:
#                     name = ent.text
#                     break

#         return {
#             "name": name,
#             "email": email.group(0) if email else None,
#             "phone": phone.group(0) if phone else None
#         }

#     # -------------------------
#     # EDUCATION (STRICT SEPARATION)
#     # -------------------------
#     def extract_education(self, text):

#         doc = self.nlp(text)

#         education = set()

#         degree_keywords = [
#             "bachelor", "master", "bsc", "msc", "mba",
#             "phd", "b.tech", "m.tech", "b.e", "m.e"
#         ]

#         university_keywords = [
#             "university", "college", "institute", "school"
#         ]

#         for ent in doc.ents:

#             t = ent.text.lower()

#             # DEGREE
#             if any(d in t for d in degree_keywords):
#                 education.add(ent.text.strip())

#             # UNIVERSITY
#             elif ent.label_ == "ORG" and any(u in t for u in university_keywords):
#                 education.add(ent.text.strip())

#         return list(education) if education else ["Not Found"]

#     # -------------------------
#     # EXPERIENCE CLEAN FILTER
#     # -------------------------
#     def is_valid_company(self, name):

#         name = name.lower().strip()

#         blacklist = [
#             "university", "college", "school",
#             "project", "management", "leadership",
#             "director", "engineer", "developer",
#             "forum", "association", "course"
#         ]

#         if any(b in name for b in blacklist):
#             return False

#         if len(name.split()) < 2:
#             return False

#         return True

#     # -------------------------
#     # EXPERIENCE EXTRACTION
#     # -------------------------
#     def extract_experience(self, text):

#         lines = [l.strip() for l in text.split("\n") if l.strip()]

#         experience = []
#         seen = set()

#         for i, line in enumerate(lines):

#             match = re.search(
#                 r"(20\d{2})\s*[-–]\s*(present|current|20\d{2})",
#                 line,
#                 re.I
#             )

#             if not match:
#                 continue

#             start, end = match.groups()

#             context = " ".join(lines[max(0, i-2): i+2])
#             doc = self.nlp(context)

#             for ent in doc.ents:

#                 if ent.label_ != "ORG":
#                     continue

#                 company = ent.text.strip()

#                 if not self.is_valid_company(company):
#                     continue

#                 key = company.lower()

#                 if key in seen:
#                     continue

#                 experience.append({
#                     "company": company,
#                     "tenure": f"{start} - {end}"
#                 })

#                 seen.add(key)

#         return experience[:6]

#     # -------------------------
#     # TOTAL EXPERIENCE (MERGE OVERLAPS)
#     # -------------------------
#     def calculate_total_experience(self, experience):

#         current_year = datetime.now().year
#         periods = []

#         for exp in experience:

#             try:
#                 tenure = exp.get("tenure", "")

#                 if "unknown" in tenure.lower():
#                     continue

#                 start, end = tenure.split("-")

#                 start = int(start.strip())
#                 end = end.strip().lower()

#                 end = current_year if "present" in end else int(end)

#                 periods.append((start, end))

#             except:
#                 continue

#         periods.sort()

#         merged = []

#         for s, e in periods:

#             if not merged:
#                 merged.append([s, e])
#                 continue

#             last = merged[-1]

#             if s <= last[1]:
#                 last[1] = max(last[1], e)
#             else:
#                 merged.append([s, e])

#         return round(sum(e - s for s, e in merged), 1)

#     # -------------------------
#     # PIPELINE
#     # -------------------------
#     def parse(self, resume_path):

#         logger.info("Processing Resume...")

#         text = self.extract_text(resume_path)

#         if not text.strip():
#             return {"error": "Empty resume"}

#         bio = self.extract_bio(text)
#         edu = self.extract_education(text)
#         exp = self.extract_experience(text)
#         total_exp = self.calculate_total_experience(exp)

#         ats = self.skill_matcher.match_skills_from_files(
#             self.jd_text,
#             text
#         )

#         return {
#             "biographical": bio,
#             "experience": exp,
#             "total_experience_years": total_exp,
#             "education": edu,
#             "ats_result": ats,
#             "final_score": ats.get("match_score", 0),
#             "decision": ats.get("decision", "REJECT"),
#             "metadata": {
#                 "file": os.path.basename(resume_path),
#                 "processed_at": datetime.now().isoformat()
#             }
#         }

##############


import os
import re
import fitz
import magic
import pytesseract
import spacy
from docx import Document

from skill_matcher import SkillMatcher


class UniversalATSParser:
    
    def __init__(self, jd_path=None):

        self.skill_matcher = SkillMatcher()
        self.nlp = spacy.load("en_core_web_sm")
        self._cache = {}

        self.jd_text = self.preprocess(jd_path) if jd_path else ""

    # =========================
    # TEXT LOADER
    # =========================
    def extract_text(self, path):

        if not path or not os.path.exists(path):
            return ""

        if path in self._cache:
            return self._cache[path]

        mime = magic.Magic(mime=True).from_file(path)

        if "pdf" in mime:
            doc = fitz.open(path)
            text = "\n".join(p.get_text() for p in doc)

        elif "word" in mime or "docx" in mime:
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs)

        elif "image" in mime:
            text = pytesseract.image_to_string(path)

        else:
            with open(path, "r", errors="ignore") as f:
                text = f.read()

        self._cache[path] = text
        return text

    # =========================
    # CLEAN OCR
    # =========================
    def clean_text(self, text):
        lines = text.split("\n")
        cleaned = []

        for l in lines:
            l = re.sub(r"\s+", " ", l).strip()
            if len(l) > 1:
                cleaned.append(l)

        return "\n".join(cleaned)

    def preprocess(self, path):
        return self.clean_text(self.extract_text(path))

    # =========================
    # SECTION SPLITTER (IMPORTANT FIX)
    # =========================
    def split_sections(self, text):

        sections = {
            "experience": [],
            "education": [],
            "skills": [],
            "other": []
        }

        current = "other"

        for line in text.split("\n"):
            low = line.lower()

            if "experience" in low or "work" in low:
                current = "experience"
                continue

            if "education" in low:
                current = "education"
                continue

            if "skill" in low:
                current = "skills"
                continue

            sections[current].append(line)

        return sections

    # =========================
    # NAME (FIXED HARD)
    # =========================
    def extract_name(self, text, doc):

        # spaCy PERSON first
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                name = ent.text.strip()
                if len(name.split()) <= 4:
                    return name

        # fallback: top clean lines
        lines = text.split("\n")[:8]

        blacklist = ["resume", "email", "phone", "address", "india"]

        for line in lines:
            line = line.strip()

            if any(b in line.lower() for b in blacklist):
                continue

            if 2 <= len(line.split()) <= 3 and line.replace(" ", "").isalpha():
                return line

        return "Unknown"

    # =========================
    # EMAIL
    # =========================
    def extract_email(self, text):
        m = re.findall(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
        return m[0] if m else None

    # =========================
    # PHONE (FIXED)
    # =========================
    def extract_phone(self, text):

        m = re.findall(r"\+?\d[\d\s\-]{8,}", text)

        for p in m:
            p_clean = re.sub(r"\D", "", p)

            if 10 <= len(p_clean) <= 13:
                return p

        return None

    # =========================
    # EXPERIENCE (CLEANED LOGIC)
    # =========================
    def extract_experience(self, text):

        experience = []

        lines = text.split("\n")

        for line in lines:

            if re.search(r"(education|university|college|institute)", line, re.I):
                continue  # IMPORTANT FIX

            years = re.search(r"(20\d{2}|19\d{2})\s*(?:-|to|–)\s*(present|20\d{2}|19\d{2})", line, re.I)

            if not years:
                continue

            words = re.findall(r"[A-Z][A-Za-z&]{2,}", line)

            if not words:
                continue

            company = words[0]

            if company.lower() in ["aug", "jul", "jan", "feb"]:
                continue

            experience.append({
                "company": company,
                "role": "Unknown",
                "tenure": f"{years.group(1)} - {years.group(2)}"
            })

        return experience[:6]

    # =========================
    # EDUCATION (CLEAN FIX)
    # =========================
    def extract_education(self, text):

        education = []

        lines = text.split("\n")

        for line in lines:

            if not re.search(r"(university|institute|college|b\.?tech|m\.?tech|bsc|msc)", line, re.I):
                continue

            education.append({
                "institution": line[:120],
                "years": "Unknown"
            })

        return education

    # =========================
    # ATS SCORE (STABLE)
    # =========================
    def compute_score(self, resume_skills, jd_skills):

        if not jd_skills:
            return 0

        matched = set(resume_skills) & set(jd_skills)

        coverage = len(matched) / len(jd_skills)

        score = (coverage * 85) + (len(matched) * 2)

        return min(round(score, 2), 100)

    # =========================
    # MAIN PIPELINE
    # =========================
    def parse(self, resume_path):

        text = self.preprocess(resume_path)
        doc = self.nlp(text)

        name = self.extract_name(text, doc)
        email = self.extract_email(text)
        phone = self.extract_phone(text)

        experience = self.extract_experience(text)
        education = self.extract_education(text)

        resume_skills = self.skill_matcher.extract_skills(text)
        jd_skills = self.skill_matcher.extract_skills(self.jd_text)

        matched = list(set(resume_skills) & set(jd_skills))
        missing = list(set(jd_skills) - set(resume_skills))

        score = self.compute_score(resume_skills, jd_skills)

        return {
            "biographical": {
                "name": name,
                "email": email,
                "phone": phone
            },
            "experience": experience,
            "education": education,
            "skills": {
                "jd_skills": jd_skills,
                "resume_skills": resume_skills,
                "matched": matched,
                "missing": missing
            },
            "ats": {
                "score": score,
                "decision": "SHORTLIST" if score >= 50 else "REJECT"
            }
        }